#!/usr/bin/env python3
"""
AI Detection Tool - Compares student submissions with responses from different LLMs
to detect potentially AI-generated content.
"""

import os
import sys
import argparse
import json
import hashlib
import re
import time
from datetime import datetime
from pathlib import Path
import docx
import openpyxl
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter
import pandas as pd
import numpy as np
import PyPDF2

# Import NLTK with SSL fix
import nltk
nltk.download('punkt', quiet=True)
nltk.download('stopwords', quiet=True)
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords

from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from tqdm import tqdm
from dotenv import load_dotenv

# Fix NLTK SSL certificate issue
import ssl

# More robust NLTK data download
try:
    nltk.download('punkt')
    nltk.download('stopwords')
    # Test if punkt is working
    from nltk.tokenize import sent_tokenize
    sent_tokenize("This is a test sentence.")
except Exception as e:
    print(f"Warning: NLTK download or initialization failed: {e}")
    print("Using fallback tokenizer...")
    
    # Define fallback tokenizer functions
    def sent_tokenize_fallback(text):
        """Simple sentence tokenizer as fallback."""
        # Replace common abbreviations to avoid incorrect splitting
        for abbr in ['Mr.', 'Mrs.', 'Dr.', 'Ms.', 'Prof.', 'e.g.', 'i.e.', 'etc.']:
            text = text.replace(abbr, abbr.replace('.', '<DOT>'))
        
        # Split on sentence terminators
        sentences = []
        for chunk in text.split('\n'):
            for sent in re.split(r'(?<=[.!?])\s+', chunk):
                if sent:
                    # Restore abbreviations
                    sent = sent.replace('<DOT>', '.')
                    sentences.append(sent)
        
        return sentences

    def word_tokenize_fallback(text):
        """Simple word tokenizer as fallback."""
        # Remove punctuation and split on whitespace
        text = re.sub(r'[^\w\s]', ' ', text)
        return [word for word in text.lower().split() if word]
    
    # Replace NLTK tokenizers with fallbacks
    from nltk.tokenize import sent_tokenize as nltk_sent_tokenize
    from nltk.tokenize import word_tokenize as nltk_word_tokenize
    nltk.tokenize.sent_tokenize = sent_tokenize_fallback
    nltk.tokenize.word_tokenize = word_tokenize_fallback
    sent_tokenize = sent_tokenize_fallback
    word_tokenize = word_tokenize_fallback

try:
    _create_unverified_https_context = ssl._create_unverified_context
except AttributeError:
    pass
else:
    ssl._create_default_https_context = _create_unverified_https_context

# Load environment variables
load_dotenv()

# Import API clients
from openai import OpenAI
from anthropic import Anthropic
import google.generativeai as genai

# =====================================================================
# Configuration and Setup
# =====================================================================

# Initialize API clients
try:
    openai_client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
    anthropic_client = Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
    genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
except Exception as e:
    print(f"Error initializing API clients: {e}")
    print("Please check your API keys in the .env file.")

# Define the LLMs to test
LLM_CONFIGS = {
    "openai_4o": {
        "provider": "openai",
        "model": "chatgpt-4o-latest",
        
    },
    "openai_o3": {
        "provider": "openai",
        "model": "o3-mini",
        
    },
    "openai_o1": {
        "provider": "openai",
        "model": "o1-2024-12-17",
        
    },
    "gemini_1.5": {
        "provider": "gemini",
        "model": "gemini-1.5-pro",
        
    },
    "gemini_2.0": {
        "provider": "gemini",
        "model": "gemini-2.0-flash",
        
    },
    "claude_3.5": {
        "provider": "anthropic",
        "model": "claude-3-5-sonnet-20240620",
        
    },
    "claude_3.7": {
        "provider": "anthropic",
        "model": "claude-3-7-sonnet-20250219",
        
    }
}

# =====================================================================
# LLM API Functions
# =====================================================================

def get_openai_response(prompt, model):
    """Get a response from OpenAI API."""
    try:
        response = openai_client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
        )
        return response.choices[0].message.content
    except Exception as e:
        print(f"Error with OpenAI {model}: {e}")
        return f"Error: {str(e)}"

def get_anthropic_response(prompt, model, max_tokens=2000):
    """Get a response from Anthropic API."""
    try:
        response = anthropic_client.messages.create(
            model=model,
            max_tokens=max_tokens,
            messages=[{"role": "user", "content": prompt}]
        )
        return response.content[0].text
    except Exception as e:
        print(f"Error with Anthropic {model}: {e}")
        return f"Error: {str(e)}"

def get_gemini_response(prompt, model, max_tokens=2000):
    """Get a response from Gemini API."""
    try:
        gemini_model = genai.GenerativeModel(model_name=model)
        response = gemini_model.generate_content(prompt, generation_config={"max_output_tokens": max_tokens})
        return response.text
    except Exception as e:
        print(f"Error with Gemini {model}: {e}")
        return f"Error: {str(e)}"

def get_all_llm_responses(prompt, cache=True):
    """Get responses from all configured LLMs with optional caching."""
    results = {}
    
    # Create cache directory if it doesn't exist
    cache_dir = "llm_response_cache"
    os.makedirs(cache_dir, exist_ok=True)
    
    # Generate a cache key based on the prompt
    cache_key = hashlib.md5(prompt.encode()).hexdigest()
    cache_file = os.path.join(cache_dir, f"{cache_key}.json")
    
    # Try to load from cache first if available and cache is enabled
    if cache and os.path.exists(cache_file):
        try:
            with open(cache_file, 'r') as f:
                cached_results = json.load(f)
                print("Loaded LLM responses from cache.")
                return cached_results
        except Exception as e:
            print(f"Error loading cache: {e}")
    
    # If not using cache or cache failed, get fresh responses
    for llm_name, config in tqdm(LLM_CONFIGS.items(), desc="Getting LLM responses"):
        try:
            provider = config["provider"]
            model = config["model"]
            max_tokens = config["max_tokens"]
            
            if provider == "openai":
                response = get_openai_response(prompt, model, max_tokens)
            elif provider == "anthropic":
                response = get_anthropic_response(prompt, model, max_tokens)
            elif provider == "gemini":
                response = get_gemini_response(prompt, model, max_tokens)
            else:
                response = "Unsupported provider"
            
            results[llm_name] = response
            
            # Add a small delay to avoid rate limits
            time.sleep(1)
            
        except Exception as e:
            results[llm_name] = f"Error: {str(e)}"
    
    # Save responses to cache if enabled
    if cache:
        try:
            with open(cache_file, 'w') as f:
                json.dump(results, f)
        except Exception as e:
            print(f"Error saving to cache: {e}")
    
    return results

# =====================================================================
# Text Analysis Functions
# =====================================================================

# def compute_text_metrics(text):
#     """Compute various text metrics for analysis."""
#     if not text or text.startswith("Error:"):
#         return {
#             "avg_sentence_length": 0,
#             "avg_word_length": 0,
#             "sentence_count": 0,
#             "word_count": 0,
#             "unique_words_ratio": 0,
#             "complexity_score": 0
#         }
    
#     # Clean text and tokenize
#     text = re.sub(r'\s+', ' ', text).strip()
#     sentences = sent_tokenize(text)
#     words = word_tokenize(text.lower())
    
#     # Filter out punctuation
#     words = [word for word in words if word.isalnum()]
    
#     # Basic metrics
#     sentence_count = len(sentences)
#     word_count = len(words)
#     unique_words = set(words)
#     unique_words_ratio = len(unique_words) / word_count if word_count > 0 else 0
    
#     # Average lengths
#     avg_sentence_length = word_count / sentence_count if sentence_count > 0 else 0
#     avg_word_length = sum(len(word) for word in words) / word_count if word_count > 0 else 0
    
#     # Complexity score (simple heuristic)
#     complexity_score = (avg_sentence_length * 0.5) + (avg_word_length * 1.5) + (unique_words_ratio * 5)
    
#     return {
#         "avg_sentence_length": avg_sentence_length,
#         "avg_word_length": avg_word_length,
#         "sentence_count": sentence_count,
#         "word_count": word_count,
#         "unique_words_ratio": unique_words_ratio,
#         "complexity_score": complexity_score
#     }


def compute_text_metrics(text):
    """Compute various text metrics for analysis."""
    if not text or (isinstance(text, str) and text.startswith("Error:")):
        return {
            "avg_sentence_length": 0,
            "avg_word_length": 0,
            "sentence_count": 0,
            "word_count": 0,
            "unique_words_ratio": 0,
            "complexity_score": 0
        }
    
    try:
        # Clean text and tokenize
        text = re.sub(r'\s+', ' ', str(text)).strip()
        
        # Use try-except to handle tokenization failures
        try:
            sentences = sent_tokenize(text)
        except Exception:
            # Super basic fallback
            sentences = text.split('. ')
        
        try:
            words = word_tokenize(text.lower())
        except Exception:
            # Super basic fallback
            words = text.lower().split()
        
        # Filter out punctuation
        words = [word for word in words if word.isalnum()]
        
        # Basic metrics
        sentence_count = len(sentences)
        word_count = len(words)
        unique_words = set(words)
        unique_words_ratio = len(unique_words) / word_count if word_count > 0 else 0
        
        # Average lengths
        avg_sentence_length = word_count / sentence_count if sentence_count > 0 else 0
        avg_word_length = sum(len(word) for word in words) / word_count if word_count > 0 else 0
        
        # Complexity score (simple heuristic)
        complexity_score = (avg_sentence_length * 0.5) + (avg_word_length * 1.5) + (unique_words_ratio * 5)
        
        return {
            "avg_sentence_length": avg_sentence_length,
            "avg_word_length": avg_word_length,
            "sentence_count": sentence_count,
            "word_count": word_count,
            "unique_words_ratio": unique_words_ratio,
            "complexity_score": complexity_score
        }
    except Exception as e:
        print(f"Error computing text metrics: {e}")
        # Return default metrics on error
        return {
            "avg_sentence_length": 0,
            "avg_word_length": 0,
            "sentence_count": 0,
            "word_count": 0,
            "unique_words_ratio": 0,
            "complexity_score": 0
        }

def compare_responses(llm_responses, student_answer):
    """Compare each LLM response to the student answer."""
    results = {}
    
    # Calculate metrics for student answer
    student_metrics = compute_text_metrics(student_answer)
    
    # Create a list of all texts for TF-IDF vectorization
    all_texts = [student_answer]
    llm_names = []
    
    for llm_name, response in llm_responses.items():
        if not response.startswith("Error:"):
            all_texts.append(response)
            llm_names.append(llm_name)
            
            # Calculate metrics for this LLM response
            llm_metrics = compute_text_metrics(response)
            
            # Calculate metric differences
            metric_diffs = {
                f"{k}_diff": abs(v - student_metrics[k])
                for k, v in llm_metrics.items()
            }
            
            results[llm_name] = {
                "student_metrics": student_metrics,
                "llm_metrics": llm_metrics,
                "metric_diffs": metric_diffs,
            }
    
    # Calculate TF-IDF similarity if we have valid responses
    if len(all_texts) > 1:
        vectorizer = TfidfVectorizer(stop_words='english')
        tfidf_matrix = vectorizer.fit_transform(all_texts)
        
        # Calculate cosine similarity between student answer and each LLM response
        for i, llm_name in enumerate(llm_names):
            similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[i+1:i+2])[0][0]
            results[llm_name]["similarity_score"] = similarity
    
    return results

def detect_ai_patterns(text):
    """
    Detect patterns commonly associated with AI-generated text.
    Returns a dictionary of detected patterns and their scores.
    """
    patterns = {
        'repetitive_phrases': 0,
        'formulaic_structure': 0,
        'conventional_transitions': 0,
        'generic_examples': 0,
        'overuse_of_hedging': 0
    }
    
    # Check for repetitive phrases
    words = word_tokenize(text.lower())
    word_pairs = [' '.join(words[i:i+2]) for i in range(len(words)-1)]
    word_triplets = [' '.join(words[i:i+3]) for i in range(len(words)-2)]
    
    # Count repeated phrases
    pair_counts = {}
    for pair in word_pairs:
        if pair not in pair_counts:
            pair_counts[pair] = 0
        pair_counts[pair] += 1
    
    triplet_counts = {}
    for triplet in word_triplets:
        if triplet not in triplet_counts:
            triplet_counts[triplet] = 0
        triplet_counts[triplet] += 1
    
    # Score for repetitiveness (higher is more repetitive)
    repeated_pairs = sum(1 for count in pair_counts.values() if count > 3)
    repeated_triplets = sum(1 for count in triplet_counts.values() if count > 2)
    
    patterns['repetitive_phrases'] = min(1.0, (repeated_pairs * 0.1 + repeated_triplets * 0.2))
    
    # Check for formulaic structure
    sentences = sent_tokenize(text)
    
    # Look for conventional AI patterns in introductions and conclusions
    intro_patterns = [
        r'(?i)in this (essay|paper|response|assignment)',
        r'(?i)i (will|am going to) (discuss|explore|examine|analyze)',
        r'(?i)(firstly|first of all|to begin with)',
        r'(?i)this (essay|paper|response) (aims|intends) to'
    ]
    
    concl_patterns = [
        r'(?i)in conclusion',
        r'(?i)to (sum|summarize) up',
        r'(?i)as (mentioned|discussed|shown) above',
        r'(?i)in summary'
    ]
    
    intro_matches = 0
    for pattern in intro_patterns:
        if re.search(pattern, sentences[0] if sentences else ""):
            intro_matches += 1
    
    concl_matches = 0
    for pattern in concl_patterns:
        if sentences and re.search(pattern, sentences[-1]):
            concl_matches += 1
    
    patterns['formulaic_structure'] = min(1.0, (intro_matches * 0.25 + concl_matches * 0.25))
    
    # Check for conventional transitions
    transition_words = [
        r'(?i)\b(however|nevertheless|furthermore|moreover|in addition|consequently)\b',
        r'(?i)\b(therefore|thus|hence|as a result|accordingly)\b',
        r'(?i)\b(for instance|for example|specifically|in particular)\b',
        r'(?i)\b(firstly|secondly|thirdly|finally)\b'
    ]
    
    transition_count = 0
    for pattern in transition_words:
        transition_count += len(re.findall(pattern, text))
    
    patterns['conventional_transitions'] = min(1.0, transition_count / (len(sentences) * 0.8) if sentences else 0)
    
    # Check for generic examples
    generic_examples = [
        r'(?i)for (example|instance)',
        r'(?i)such as',
        r'(?i)in other words',
        r'(?i)to illustrate'
    ]
    
    example_count = 0
    for pattern in generic_examples:
        example_count += len(re.findall(pattern, text))
    
    patterns['generic_examples'] = min(1.0, example_count / (len(sentences) * 0.5) if sentences else 0)
    
    # Check for overuse of hedging language
    hedging_words = [
        r'(?i)\b(may|might|could|possibly|potentially|perhaps|probably)\b',
        r'(?i)\b(generally|usually|typically|often|sometimes|frequently)\b',
        r'(?i)\b(tend to|seem to|appear to)\b',
        r'(?i)\b(relatively|fairly|quite|rather|somewhat)\b'
    ]
    
    hedging_count = 0
    for pattern in hedging_words:
        hedging_count += len(re.findall(pattern, text))
    
    patterns['overuse_of_hedging'] = min(1.0, hedging_count / (len(sentences) * 0.7) if sentences else 0)
    
    # Calculate overall AI pattern score (0 to 1 scale)
    overall_score = (
        patterns['repetitive_phrases'] * 0.2 +
        patterns['formulaic_structure'] * 0.3 +
        patterns['conventional_transitions'] * 0.15 +
        patterns['generic_examples'] * 0.15 +
        patterns['overuse_of_hedging'] * 0.2
    )
    
    patterns['overall_pattern_score'] = overall_score
    
    return patterns

def calculate_ai_probability(comparison_results):
    """Calculate probability that student answer was AI-generated."""
    probabilities = {}
    
    # If we have no valid results, add a default low probability
    if not comparison_results:
        return {"fallback_model": 0.1}  # Return a default low probability
    
    for llm_name, results in comparison_results.items():
        if "similarity_score" not in results:
            probabilities[llm_name] = 0
            continue
            
        # Factors that contribute to AI probability:
        # 1. High text similarity (30%)
        # 2. Similar complexity scores (25%)
        # 3. Similar sentence length patterns (15%)
        # 4. Similar unique word ratios (20%)
        # 5. Similar word length patterns (10%)
        
        sim_score = results["similarity_score"]
        complexity_diff = results["metric_diffs"]["complexity_score_diff"]
        max_complexity = max(results["student_metrics"]["complexity_score"], 
                            results["llm_metrics"]["complexity_score"])
        
        # Normalize complexity difference
        norm_complexity_diff = 1 - (complexity_diff / max_complexity if max_complexity > 0 else 1)
        norm_complexity_diff = max(0, min(1, norm_complexity_diff))
        
        # Sentence length similarity
        sent_len_diff = results["metric_diffs"]["avg_sentence_length_diff"]
        max_sent_len = max(results["student_metrics"]["avg_sentence_length"],
                          results["llm_metrics"]["avg_sentence_length"])
        norm_sent_len_diff = 1 - (sent_len_diff / max_sent_len if max_sent_len > 0 else 1)
        norm_sent_len_diff = max(0, min(1, norm_sent_len_diff))
        
        # Word ratio similarity
        word_ratio_diff = results["metric_diffs"]["unique_words_ratio_diff"]
        norm_word_ratio_diff = 1 - word_ratio_diff
        norm_word_ratio_diff = max(0, min(1, norm_word_ratio_diff))
        
        # Word length similarity
        word_len_diff = results["metric_diffs"]["avg_word_length_diff"]
        max_word_len = max(results["student_metrics"]["avg_word_length"],
                          results["llm_metrics"]["avg_word_length"])
        norm_word_len_diff = 1 - (word_len_diff / max_word_len if max_word_len > 0 else 1)
        norm_word_len_diff = max(0, min(1, norm_word_len_diff))
        
        # Weighted probability calculation
        probability = (
            0.30 * sim_score +
            0.25 * norm_complexity_diff +
            0.15 * norm_sent_len_diff +
            0.20 * norm_word_ratio_diff +
            0.10 * norm_word_len_diff
        )
        
        probabilities[llm_name] = probability
    
    return probabilities

def normalize_text(text):
    """Normalize text by removing excessive whitespace and standardizing formatting."""
    # Remove excessive newlines and spaces
    text = re.sub(r'\n+', '\n', text)
    text = re.sub(r' +', ' ', text)
    text = text.strip()
    
    # Try to detect and normalize list formats (numbered, bullet points, etc.)
    text = re.sub(r'^\s*(\d+)[.)\]]\s*', r'\1. ', text, flags=re.MULTILINE)
    text = re.sub(r'^\s*[-•*]\s*', '- ', text, flags=re.MULTILINE)
    
    return text

# =====================================================================
# Document Processing Functions
# =====================================================================

def load_questions(file_path):
    """Load questions from a file."""
    questions = []
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            for line in file:
                line = line.strip()
                if line:  # Skip empty lines
                    questions.append(line)
        return questions
    except Exception as e:
        print(f"Error loading questions: {e}")
        return []
    
def load_text_from_pdf(file_path):
    """Extract text from a PDF file."""
    try:
        text = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text.append(page.extract_text())
        return '\n'.join(text)
    except Exception as e:
        print(f"Error extracting text from PDF {file_path}: {e}")
        return ""
    
def load_student_responses(directory_path):
    """Load student responses from Word DOCX and PDF files in a directory."""
    student_responses = {}
    try:
        for filename in os.listdir(directory_path):
            file_path = os.path.join(directory_path, filename)
            
            # Extract student ID with consistent format (including LATE if present)
            student_name_parts = os.path.splitext(filename)[0].split('_')
            if len(student_name_parts) >= 3:
                if 'LATE' in student_name_parts:
                    # Include LATE in the student ID if it exists
                    late_index = student_name_parts.index('LATE')
                    student_id = '_'.join(student_name_parts[:late_index+3])  # Include name, LATE, and numbers
                else:
                    # Regular format - include name and numbers
                    student_id = '_'.join(student_name_parts[:3])
            else:
                # Fallback for unusual filename formats
                student_id = os.path.splitext(filename)[0]
            
            # Process based on file type
            if filename.endswith('.docx'):
                # Extract text from DOCX file
                doc = docx.Document(file_path)
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
                student_responses[student_id] = '\n'.join(full_text)
            elif filename.endswith('.pdf'):
                # Extract text from PDF file
                student_responses[student_id] = load_text_from_pdf(file_path)
        
        print(f"Loaded {len(student_responses)} student submissions")
        return student_responses
    except Exception as e:
        print(f"Error loading student responses: {e}")
        return {}

def load_student_responses_docx(directory_path):
    """Load student responses from Word DOCX files in a directory."""
    student_responses = {}
    try:
        for filename in os.listdir(directory_path):
            if filename.endswith('.docx'):
                file_path = os.path.join(directory_path, filename)
                student_name_parts = os.path.splitext(filename)[0].split('_')
                if len(student_name_parts) >= 3:
                    if 'LATE' in student_name_parts:
                        # Include LATE in the student ID if it exists
                        late_index = student_name_parts.index('LATE')
                        student_id = '_'.join(student_name_parts[:late_index+3])  # Include name, LATE, and numbers
                    else:
                        # Regular format - include name and numbers
                        student_id = '_'.join(student_name_parts[:3])
                else:
                    # Fallback for unusual filename formats
                    student_id = os.path.splitext(filename)[0]
                
                # Extract text from DOCX file
                doc = docx.Document(file_path)
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
                
                student_responses[student_id] = '\n'.join(full_text)
        
        print(f"Loaded {len(student_responses)} DOCX student submissions")
        return student_responses
    except Exception as e:
        print(f"Error loading student responses: {e}")
        return {}

def extract_answer_for_question(full_response, question, question_idx):
    """
    Extract the answer for a specific question from the full response text.
    Handles various document formats and structures.
    """
    # Strategy 1: Look for question numbers (Q1, Question 1, etc.)
    patterns = [
        rf"(?:^|\n)(?:Question|Q)\s*{question_idx+1}(?::|\.|\)|\s)(.*?)(?=(?:^|\n)(?:Question|Q)\s*{question_idx+2}(?::|\.|\)|\s)|$)",
        rf"(?:^|\n){question_idx+1}(?:\.|\)|\s)(.*?)(?=(?:^|\n){question_idx+2}(?:\.|\)|\s)|$)",
        rf"(?:^|\n)(?:Problem|Exercise)\s*{question_idx+1}(?::|\.|\)|\s)(.*?)(?=(?:^|\n)(?:Problem|Exercise)\s*{question_idx+2}(?::|\.|\)|\s)|$)"
    ]
    
    for pattern in patterns:
        match = re.search(pattern, full_response, re.DOTALL)
        if match:
            return match.group(1).strip()
    
    # Strategy 2: Look for the question text
    if len(question) > 10:
        # Get first ~50 chars of question as search pattern
        search_text = question[:min(50, len(question))]
        search_text = re.escape(search_text)  # Escape special regex characters
        
        # Try to find the question followed by an answer
        match = re.search(rf"{search_text}.*?(?:\n|$)(.*?)(?=(?:Question|Q)\s*\d|$)", 
                          full_response, re.DOTALL | re.IGNORECASE)
        
        if match:
            return match.group(1).strip()
    
    # Strategy 3: Look for section breaks
    sections = re.split(r'\n{2,}', full_response)
    
    # Try to identify sections that look like questions and answers
    question_sections = []
    for i, section in enumerate(sections):
        if re.search(r'^(?:Question|Q)\s*\d+|^\d+[\.)\s]', section.strip()):
            question_sections.append(i)
    
    # If we found question markers, extract the content between them
    if question_sections and len(question_sections) > question_idx:
        start_idx = question_sections[question_idx]
        end_idx = question_sections[question_idx + 1] if question_idx + 1 < len(question_sections) else len(sections)
        
        # Combine the sections between the question markers
        answer_sections = sections[start_idx:end_idx]
        # Skip the first section if it contains the question
        if re.search(r'^(?:Question|Q)\s*\d+|^\d+[\.)\s]', answer_sections[0].strip()):
            answer_sections = answer_sections[1:]
        
        if answer_sections:
            return '\n'.join(answer_sections).strip()
    
    # Strategy 4: If we have a reasonable number of sections, assume they correspond to questions
    if 3 <= len(sections) <= 20:  # Reasonable range for number of sections
        # Assume each section is a separate question/answer
        section_idx = question_idx
        if section_idx < len(sections):
            return sections[section_idx].strip()
    
    # Strategy 5: As a fallback, divide the document into equal parts
    num_questions = 5  # Assume 5 questions by default, adjust as needed
    segment_size = len(full_response) // num_questions
    start_pos = question_idx * segment_size
    end_pos = (question_idx + 1) * segment_size if question_idx < num_questions - 1 else len(full_response)
    
    return full_response[start_pos:end_pos].strip()

# =====================================================================
# Analysis and Reporting Functions
# =====================================================================

def analyze_student_answer(prompt, student_answer, cache=True, display_responses=False):
    """Run the full analysis pipeline on a student answer."""
    print("Fetching LLM responses...")
    llm_responses = get_all_llm_responses(prompt, cache=cache)
    
    if display_responses:
        print("\n=== LLM Responses ===")
        for llm, response in llm_responses.items():
            print(f"\n**{llm}**:")
            print("-" * 80)
            print(response[:500] + "..." if len(response) > 500 else response)
            print("-" * 80)
    
    print("\nComparing responses...")
    comparison_results = compare_responses(llm_responses, student_answer)
    
    print("\nCalculating AI generation probabilities...")
    ai_probabilities = calculate_ai_probability(comparison_results)
    
    # Detect AI-specific patterns
    pattern_results = detect_ai_patterns(student_answer)
    
    # Combine pattern detection with similarity analysis
    # We'll use a 70% weight on similarity-based detection and 30% on pattern-based detection
    combined_probabilities = {}
    for llm_name, prob in ai_probabilities.items():
        combined_probabilities[llm_name] = 0.7 * prob + 0.3 * pattern_results['overall_pattern_score']
    
    # Get the LLM with highest probability
    if combined_probabilities:
        max_llm = max(combined_probabilities.items(), key=lambda x: x[1])
        overall_ai_score = max_llm[1]
        most_likely_model = max_llm[0] if overall_ai_score > 0.5 else "Human-written"
        
    else:
        # Default values if no probabilities were calculated
        overall_ai_score = 0.1
        max_llm = ("unknown_model", overall_ai_score)
        most_likely_model = "Analysis-failed"
    
    # Generate a summary report
    report = {
        "prompt": prompt,
        "student_answer_preview": student_answer[:150] + "..." if len(student_answer) > 150 else student_answer,
        "overall_ai_score": overall_ai_score,
        # "most_likely_model": max_llm[0] if overall_ai_score > 0.5 else "Human-written",
        "most_likely_model": most_likely_model,
        "confidence": "High" if overall_ai_score > 0.8 or overall_ai_score < 0.3 else "Medium" if overall_ai_score > 0.6 or overall_ai_score < 0.4 else "Low",
        "model_probabilities": combined_probabilities,
        "pattern_analysis": pattern_results
    }
    
    print("\n=== Analysis Report ===")
    print(f"Overall AI Score: {report['overall_ai_score']:.2f}")
    print(f"Most Likely Source: {report['most_likely_model']}")
    print(f"Confidence: {report['confidence']}")
    
    return report, llm_responses, comparison_results, combined_probabilities

def generate_excel_report(results_df, output_file):
    """Generate an Excel report of AI detection results."""
    # Create workbook and sheets
    wb = openpyxl.Workbook()
    
    # Summary sheet
    summary_sheet = wb.active
    summary_sheet.title = "Summary"
    
    # Add title and headers
    summary_sheet['A1'] = "AI-Generated Content Detection Report"
    summary_sheet['A1'].font = Font(size=14, bold=True)
    summary_sheet.merge_cells('A1:G1')
    summary_sheet['A1'].alignment = Alignment(horizontal='center')
    
    # Add summary statistics
    total_submissions = len(results_df)
    likely_ai = results_df[results_df['ai_score'] > 0.7].shape[0]
    possibly_ai = results_df[(results_df['ai_score'] > 0.5) & (results_df['ai_score'] <= 0.7)].shape[0]
    likely_human = results_df[results_df['ai_score'] <= 0.5].shape[0]
    
    summary_sheet['A3'] = "Total Submissions:"
    summary_sheet['B3'] = total_submissions
    
    summary_sheet['A4'] = "Likely AI-Generated:"
    summary_sheet['B4'] = likely_ai
    summary_sheet['C4'] = f"{(likely_ai/total_submissions*100):.1f}%"
    
    summary_sheet['A5'] = "Possibly AI-Generated:"
    summary_sheet['B5'] = possibly_ai
    summary_sheet['C5'] = f"{(possibly_ai/total_submissions*100):.1f}%"
    
    summary_sheet['A6'] = "Likely Human-Written:"
    summary_sheet['B6'] = likely_human
    summary_sheet['C6'] = f"{(likely_human/total_submissions*100):.1f}%"
    
    # Style summary section
    for row in range(3, 7):
        summary_sheet[f'A{row}'].font = Font(bold=True)
    
    # Add results by question
    summary_sheet['A8'] = "Results by Question"
    summary_sheet['A8'].font = Font(size=12, bold=True)
    
    headers = ["Question ID", "Question", "Submissions", "Average AI Score", "Min Score", "Max Score"]
    for col, header in enumerate(headers, start=1):
        cell = summary_sheet.cell(row=9, column=col)
        cell.value = header
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal='center')
        cell.border = Border(bottom=Side(style='medium'))
    
    # Group data by question
    by_question = results_df.groupby('question_id').agg({
        'ai_score': ['mean', 'min', 'max', 'count'],
        'question': 'first'
    })
    
    # Add question data
    for i, (idx, row) in enumerate(by_question.iterrows(), start=10):
        summary_sheet.cell(row=i, column=1).value = idx
        summary_sheet.cell(row=i, column=2).value = row[('question', 'first')]
        summary_sheet.cell(row=i, column=3).value = row[('ai_score', 'count')]
        summary_sheet.cell(row=i, column=4).value = row[('ai_score', 'mean')]
        summary_sheet.cell(row=i, column=5).value = row[('ai_score', 'min')]
        summary_sheet.cell(row=i, column=6).value = row[('ai_score', 'max')]
    
    # Add results by student
    start_row = len(by_question) + 12
    summary_sheet.cell(row=start_row, column=1).value = "Results by Student"
    summary_sheet.cell(row=start_row, column=1).font = Font(size=12, bold=True)
    
    headers = ["Student ID", "Submissions", "Average AI Score", "Min Score", "Max Score"]
    for col, header in enumerate(headers, start=1):
        cell = summary_sheet.cell(row=start_row+1, column=col)
        cell.value = header
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal='center')
        cell.border = Border(bottom=Side(style='medium'))
    
    # Group data by student
    by_student = results_df.groupby('student_id').agg({
        'ai_score': ['mean', 'min', 'max', 'count']
    })
    
    # Add student data with conditional formatting
    red_fill = PatternFill(start_color="FFCCCC", end_color="FFCCCC", fill_type="solid")
    yellow_fill = PatternFill(start_color="FFFFCC", end_color="FFFFCC", fill_type="solid")
    green_fill = PatternFill(start_color="CCFFCC", end_color="CCFFCC", fill_type="solid")
    
    for i, (idx, row) in enumerate(by_student.iterrows(), start=start_row+2):
        avg_score = row[('ai_score', 'mean')]
        
        # Apply conditional formatting
        fill = red_fill if avg_score > 0.7 else yellow_fill if avg_score > 0.5 else green_fill
        for col in range(1, 6):
            summary_sheet.cell(row=i, column=col).fill = fill
        
        summary_sheet.cell(row=i, column=1).value = idx
        summary_sheet.cell(row=i, column=2).value = row[('ai_score', 'count')]
        summary_sheet.cell(row=i, column=3).value = avg_score
        summary_sheet.cell(row=i, column=4).value = row[('ai_score', 'min')]
        summary_sheet.cell(row=i, column=5).value = row[('ai_score', 'max')]
    
    # Create detailed results sheet
    details_sheet = wb.create_sheet(title="Detailed Results")
    
    # Add headers
    headers = ["Student ID", "Question ID", "Question", "AI Score", "Most Likely Source", "Confidence"]
    for col, header in enumerate(headers, start=1):
        cell = details_sheet.cell(row=1, column=col)
        cell.value = header
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal='center')
        cell.border = Border(bottom=Side(style='medium'))
    
    # Add all results with conditional formatting
    for i, (_, row) in enumerate(results_df.iterrows(), start=2):
        ai_score = row['ai_score']
        fill = red_fill if ai_score > 0.7 else yellow_fill if ai_score > 0.5 else green_fill
        
        for col in range(1, 7):
            details_sheet.cell(row=i, column=col).fill = fill
        
        details_sheet.cell(row=i, column=1).value = row['student_id']
        details_sheet.cell(row=i, column=2).value = row['question_id']
        details_sheet.cell(row=i, column=3).value = row['question']
        details_sheet.cell(row=i, column=4).value = ai_score
        details_sheet.cell(row=i, column=5).value = row['most_likely_source']
        details_sheet.cell(row=i, column=6).value = row['confidence']
    
    # Adjust column widths
    for sheet in [summary_sheet, details_sheet]:
        for col in range(1, 10):
            column_letter = get_column_letter(col)
            sheet.column_dimensions[column_letter].width = 15
        
        # Make question column wider
        sheet.column_dimensions['B'].width = 40
    
    # Save workbook
    wb.save(output_file)
    print(f"Excel report generated: {output_file}")
    return output_file

def batch_analyze_submissions(questions_file, responses_directory, results_file='ai_detection_results.xlsx', cache=True):
    """Run analysis on all questions and student submissions."""
    # Load questions and student responses
    questions = load_questions(questions_file)
    student_responses = load_student_responses(responses_directory)
    
    if not questions:
        print("No questions found. Please check the questions file.")
        return
    
    if not student_responses:
        print("No student responses found. Please check the responses directory.")
        return
    
    print(f"Loaded {len(questions)} questions and {len(student_responses)} student submissions")
    
    # Prepare results DataFrame
    results_data = []
    
    # Process each question for each student
    for q_idx, question in enumerate(questions):
        print(f"\nProcessing Question {q_idx+1}/{len(questions)}: {question[:50]}...")
        
        for student_id, full_response in student_responses.items():
            print(f"  Analyzing submission from Student {student_id}")
            
            # Try to extract answer to this specific question
            student_answer = extract_answer_for_question(full_response, question, q_idx)
            
            if not student_answer:
                print(f"  No answer found for Student {student_id}, Question {q_idx+1}")
                continue
            
            # Normalize the student's answer
            student_answer = normalize_text(student_answer)
            
            # Run the analysis
            report, _, _, _ = analyze_student_answer(question, student_answer, cache=cache, display_responses=False)
            
            # Store results
            results_data.append({
                'question_id': q_idx + 1,
                'question': question[:100] + "..." if len(question) > 100 else question,
                'student_id': student_id,
                'ai_score': report['overall_ai_score'],
                'most_likely_source': report['most_likely_model'],
                'confidence': report['confidence']
            })
            
            # Save intermediate results to CSV (for backup)
            pd.DataFrame(results_data).to_csv(results_file.replace('.xlsx', '.csv'), index=False)
            
            # Print basic result
            print(f"  Result: AI Score = {report['overall_ai_score']:.2f}, Source = {report['most_likely_model']}")
    
    # Final save to Excel
    results_df = pd.DataFrame(results_data)
    
    # Generate Excel report
    generate_excel_report(results_df, results_file)
    
    return results_df

# =====================================================================
# Command Line Interface
# =====================================================================

def parse_arguments():
    """Parse command line arguments."""
    parser = argparse.ArgumentParser(
        description='Compare student answers with LLM responses to detect AI-generated content.'
    )
    parser.add_argument('--questions', required=True, 
                        help='Path to file containing questions')
    parser.add_argument('--submissions', required=True, 
                        help='Directory containing student submission DOCX files')
    parser.add_argument('--output', default='ai_detection_results.xlsx', 
                        help='Output Excel file for results')
    parser.add_argument('--no-cache', action='store_true', 
                        help='Don\'t use cached LLM responses')
    parser.add_argument('--verbose', action='store_true', 
                        help='Print detailed progress information')
    return parser.parse_args()

def main():
    """Main entry point for the script."""
    # Parse command line arguments
    args = parse_arguments()
    
    # Verify input files
    if not os.path.exists(args.questions):
        print(f"Error: Questions file '{args.questions}' not found.")
        sys.exit(1)
    
    if not os.path.exists(args.submissions):
        print(f"Error: Submissions directory '{args.submissions}' not found.")
        sys.exit(1)
    
    # Check for Word DOCX and PDF files in submissions directory
    docx_files = [f for f in os.listdir(args.submissions) if f.endswith('.docx')]
    pdf_files = [f for f in os.listdir(args.submissions) if f.endswith('.pdf')]
    if not docx_files and not pdf_files:
        print(f"Warning: No .docx or .pdf files found in '{args.submissions}' directory.")
    else:
        print(f"Found {len(docx_files)} DOCX files and {len(pdf_files)} PDF files for analysis.")
    
    # Create output directory if it doesn't exist
    os.makedirs(os.path.dirname(os.path.abspath(args.output)), exist_ok=True)
    
    # Run batch analysis
    start_time = time.time()
    print(f"Starting analysis at {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    batch_analyze_submissions(
        args.questions, 
        args.submissions, 
        args.output,
        cache=not args.no_cache
    )
    
    elapsed_time = time.time() - start_time
    print(f"\nAnalysis complete! Results saved to {args.output}")
    print(f"Total processing time: {elapsed_time:.1f} seconds")


if __name__ == "__main__":
    main()